from docx import Document

def merge_word_files(file1, file2, output_file):
	# Load the first document
	doc1 = Document(file1)
	# Load the second document
	doc2 = Document(file2)
	
	doc3 = Document(file2)
	number_of_paragraphs = count_paragraphs(doc2)
	table = doc3.add_table(rows=0, cols=2)
	count_add = 0
	# Append the paragraphs from the second document to the first document
	for element1, element2 in zip(doc1.element.body, doc2.element.body):
		count_add+=1
		row = table.add_row().cells
		row[0]._element.append(element1)
		row[1]._element.append(element2)
  
	limit = count_paragraphs(doc3) - count_add 
	for element in doc3.element.body:
		number_of_paragraphs-=1
		if (number_of_paragraphs > limit):
			doc3.element.body.remove(element)
	# Save the merged document
	doc3.save(output_file)
	
 
def count_paragraphs(doc):
	"""Count the paragraphs in a document."""
	return len(doc.paragraphs)

# File paths for the input and output files
file1 = 'file1.docx'
file2 = 'file2.docx'
output_file = 'merged_document.docx'

# Call the merge function
merge_word_files(file1, file2, output_file)
