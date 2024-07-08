import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

def merge_word_files(file1, file2, output_file):
# Load the first document
	doc1 = Document(file1)
	# Load the second document
	doc2 = Document(file2)

	doc3 = Document(file2)
	number_of_paragraphs = count_paragraphs(doc2)
	section = doc3.sections[0]
	section.page_height = Inches(5.8)  # A5 height
	section.page_width = Inches(8.3)  # A5 width
	section.orientation = WD_ORIENT.LANDSCAPE
	table = doc3.add_table(rows=0, cols=2)
	count_add = 0
	# Append the paragraphs from the second document to the first document
	for element1, element2 in zip(doc1.element.body, doc2.element.body):
		count_add+=1
			
		row = table.add_row().cells
		row[0]._element.append(element1)
		row[1]._element.append(element2)
		if element1.tag.endswith('sectPr') or element2.tag.endswith('sectPr'):
			doc3.add_page_break()

	limit = count_paragraphs(doc3) - count_add 
	for element in doc3.element.body:
		number_of_paragraphs-=1
		if (number_of_paragraphs > limit):
			doc3.element.body.remove(element)
	# Save the merged document
	doc3.save(output_file)
	messagebox.showinfo("Success", "The files have been merged successfully!")


def count_paragraphs(doc):
    return len(doc.paragraphs)

def select_file1():
    global file1
    file1 = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])

def select_file2():
    global file2
    file2 = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])

def merge_files():
    if not file1 or not file2:
        messagebox.showerror("Error", "Please select both files.")
        return
    output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not output_file:
        return
    merge_word_files(file1, file2, output_file)

# GUI setup
root = tk.Tk()
root.title("Bilingual Document Merger")
root.geometry('320x200')  # Set the width of the GUI to 320

file1 = ''
file2 = ''

tk.Button(root, text="Select File 1", command=select_file1).pack(pady=5)
tk.Button(root, text="Select File 2", command=select_file2).pack(pady=5)
tk.Button(root, text="Run conversion", command=merge_files).pack(pady=5)

root.mainloop()
